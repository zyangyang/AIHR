"""
面试题API
"""
import os
import threading
import uuid
import json
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, Body
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.schemas import ApiResponse
from app.core.security import get_current_user
from app.core.celery_app import is_celery_available
from app.models import Resume, Prompt, LLMConfig, InterviewQuestion

router = APIRouter(prefix="/interviews", tags=["面试题"])


def _get_interview_task():
    """获取 Celery 任务，如果 Celery/Redis 不可用则返回 None"""
    if not is_celery_available():
        return None
    try:
        from app.tasks import generate_interview_task
        return generate_interview_task
    except Exception:
        return None


@router.get("", response_model=ApiResponse)
def get_interview_records(
    page: int = 1, page_size: int = 20, status: str = None,
    current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)
):
    query = db.query(InterviewQuestion).join(Resume, InterviewQuestion.resume_id == Resume.id)
    if status:
        query = query.filter(InterviewQuestion.status == status)
    total = query.count()
    interviews = query.order_by(InterviewQuestion.id.desc()).offset((page - 1) * page_size).limit(page_size).all()
    items = []
    for iq in interviews:
        resume = iq.resume
        items.append({
            "id": iq.id,
            "resume_id": resume.id,
            "name": resume.name,
            "phone": resume.phone,
            "email": resume.email,
            "job_title": resume.job.title if resume.job else "",
            "score": resume.scores[0].total_score if resume.scores else None,
            "status": iq.status,
            "generated_at": iq.generated_at.isoformat() if iq.generated_at else None,
        })
    return ApiResponse(data={"items": items, "total": total, "page": page, "page_size": page_size})


@router.post("/{resume_id}/generate", response_model=ApiResponse)
def generate_interview(
    resume_id: int, request: dict = Body(default={}), current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)
):
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")
    if resume.parse_status != "success":
        raise HTTPException(status_code=400, detail="简历尚未解析完成")

    prompt_id = request.get("prompt_id")
    llm_config_id = request.get("llm_config_id")
    if not prompt_id:
        prompt = db.query(Prompt).filter(Prompt.name == "面试题生成提示词-默认", Prompt.type == "interview").first()
        if not prompt:
            raise HTTPException(status_code=400, detail="请先在提示词管理中创建面试题提示词")
        prompt_id = prompt.id
    if not llm_config_id:
        llm_config = db.query(LLMConfig).filter(LLMConfig.is_active == True, LLMConfig.config_type == "chat").order_by(LLMConfig.id).first()
        if not llm_config:
            raise HTTPException(status_code=400, detail="请先在模型管理中配置大模型")
        llm_config_id = llm_config.id

    prompt = db.query(Prompt).filter(Prompt.id == prompt_id).first()
    if not prompt:
        raise HTTPException(status_code=404, detail="提示词不存在")
    llm_config = db.query(LLMConfig).filter(LLMConfig.id == llm_config_id).first()
    if not llm_config:
        raise HTTPException(status_code=404, detail="大模型配置不存在")

    task = _get_interview_task()
    if task:
        try:
            celery_task = task.delay(resume_id, prompt_id, llm_config_id)
            return ApiResponse(message="生成任务已提交", data={"task_id": celery_task.id})
        except Exception:
            pass

    def _run_generate_in_background():
        from app.tasks import generate_interview_task
        result = generate_interview_task.run(resume_id, prompt_id, llm_config_id)
        if "error" in result:
            print(f"生成面试题失败: {result['error']}")
        else:
            print(f"生成面试题完成: {result}")

    resume.interview_status = "generating"
    db.commit()
    threading.Thread(target=_run_generate_in_background, daemon=True).start()
    return ApiResponse(message="生成任务已提交", data={"task_id": str(uuid.uuid4())})


@router.get("/{resume_id}/status", response_model=ApiResponse)
def get_interview_status(resume_id: int, current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")
    interview = resume.interview_question
    if not interview:
        return ApiResponse(data={"status": "pending"})
    return ApiResponse(data={"status": interview.status, "questions": json.loads(interview.questions) if interview.questions else None})


@router.get("/{resume_id}/download/pdf")
def download_interview_pdf(resume_id: int, current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import platform

    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")
    interview = resume.interview_question
    if not interview or interview.status != "success":
        raise HTTPException(status_code=400, detail="面试题尚未生成")

    # 注册中文字体
    font_path = None
    if platform.system() == "Windows":
        windir = os.environ.get("WINDIR", "C:\\Windows")
        candidates = [
            Path(windir) / "Fonts" / "simhei.ttf",    # 黑体（.ttf 最可靠）
            Path(windir) / "Fonts" / "simsun.ttc",    # 宋体
            Path(windir) / "Fonts" / "msyh.ttc",      # 微软雅黑
        ]
        for p in candidates:
            if p.exists():
                font_path = str(p)
                break
    elif platform.system() == "Darwin":
        mac_candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
        for p in mac_candidates:
            if Path(p).exists():
                font_path = p
                break
    else:
        linux_candidates = [
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        ]
        for p in linux_candidates:
            if Path(p).exists():
                font_path = p
                break

    if not font_path:
        raise HTTPException(status_code=500, detail="系统未安装中文字体，无法生成PDF")

    font_name = 'ChineseFont'
    bold_font_name = 'ChineseFontBold'
    # 处理 .ttc 字体集合文件（需要指定 subfontIndex）
    if font_path.lower().endswith('.ttc'):
        pdfmetrics.registerFont(TTFont(font_name, font_path, subfontIndex=0))
        try:
            pdfmetrics.registerFont(TTFont(bold_font_name, font_path, subfontIndex=0))
            # 注册字体族映射，让 <b> 标签能找到加粗字体
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            registerFontFamily(font_name, normal=font_name, bold=bold_font_name, italic=font_name, boldItalic=bold_font_name)
        except:
            bold_font_name = font_name
    else:
        pdfmetrics.registerFont(TTFont(font_name, font_path))
        try:
            pdfmetrics.registerFont(TTFont(bold_font_name, font_path))
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            registerFontFamily(font_name, normal=font_name, bold=bold_font_name, italic=font_name, boldItalic=bold_font_name)
        except:
            bold_font_name = font_name

    # 文件命名规则：应聘岗位-简历姓名-手机号-面试题
    base_name = f"{resume.job.title}-{resume.name}-{resume.phone}-面试题"
    filename = f"{base_name}.pdf"
    output_dir = Path("downloads")
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / filename

    doc = SimpleDocTemplate(str(output_path), pagesize=A4)

    # 定义样式
    doc_title_style = ParagraphStyle(
        name='DocTitle',
        fontName='ChineseFont',
        fontSize=20,
        spaceAfter=16,
        alignment=1,
        leading=28,
    )
    section_style = ParagraphStyle(
        name='SectionTitle',
        fontName='ChineseFont',
        fontSize=14,
        spaceBefore=14,
        spaceAfter=8,
        leading=20,
    )
    normal_style = ParagraphStyle(
        name='ChineseNormal',
        fontName='ChineseFont',
        fontSize=11,
        spaceAfter=4,
        leading=16,
    )

    # 中文数字映射
    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    # 模块名称映射
    module_map = {
        'module_1': '专业能力考察',
        'module_2': '项目经验与实战',
        'module_3': '综合素质与潜力',
        'module_4': '行业认知与视野',
        'module_5': '团队协作与沟通',
    }

    story = []
    questions = json.loads(interview.questions)

    # 文档标题
    story.append(Paragraph(f"{base_name}", doc_title_style))
    story.append(Spacer(1, 10*mm))

    section_index = 0
    question_index = 0
    for module_key, questions_list in questions.items():
        # 模块标题：一、专业能力考察
        cn_section = cn_nums[section_index] if section_index < len(cn_nums) else str(section_index + 1)
        section_name = module_map.get(module_key, module_key.replace('_', ' '))
        story.append(Paragraph(f"{cn_section}、{section_name}", section_style))
        story.append(Spacer(1, 4*mm))

        for q in questions_list:
            question_index += 1
            # 题号 + 问题（加粗）
            story.append(Paragraph(f"<b>第{question_index}题：</b>{q['question']}", normal_style))
            intent = q.get('intent', '')
            if intent:
                story.append(Paragraph(f"<b>考察意图：</b>{intent}", normal_style))
            points = q.get('evaluation_points', [])
            if points:
                story.append(Paragraph("<b>评估要点：</b>", normal_style))
                for point in points:
                    story.append(Paragraph(f"  □ {point}", normal_style))
            story.append(Spacer(1, 4*mm))
        section_index += 1

    doc.build(story)
    return FileResponse(path=str(output_path), filename=filename, media_type="application/pdf")


@router.get("/{resume_id}/download/docx")
def download_interview_docx(resume_id: int, current_user: dict = Depends(get_current_user), db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")
    interview = resume.interview_question
    if not interview or interview.status != "success":
        raise HTTPException(status_code=400, detail="面试题尚未生成")

    # 文件命名规则：应聘岗位-简历姓名-手机号-面试题
    base_name = f"{resume.job.title}-{resume.name}-{resume.phone}-面试题"
    filename = f"{base_name}.docx"
    output_dir = Path("downloads")
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / filename

    doc = Document()

    # 中文数字映射
    cn_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    module_map = {
        'module_1': '专业能力考察',
        'module_2': '项目经验与实战',
        'module_3': '综合素质与潜力',
        'module_4': '行业认知与视野',
        'module_5': '团队协作与沟通',
    }

    # 文档标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(base_name)
    title_run.bold = True
    title_run.font.size = Pt(20)

    doc.add_paragraph()

    section_index = 0
    question_index = 0
    questions = json.loads(interview.questions)
    for module_key, questions_list in questions.items():
        cn_section = cn_nums[section_index] if section_index < len(cn_nums) else str(section_index + 1)
        section_name = module_map.get(module_key, module_key.replace('_', ' '))

        section_p = doc.add_paragraph()
        section_run = section_p.add_run(f"{cn_section}、{section_name}")
        section_run.bold = True
        section_run.font.size = Pt(14)

        for q in questions_list:
            question_index += 1
            q_p = doc.add_paragraph()
            q_label = q_p.add_run(f"第{question_index}题：")
            q_label.bold = True
            q_label.font.size = Pt(12)
            q_text = q_p.add_run(q['question'])
            q_text.font.size = Pt(12)

            intent = q.get('intent', '')
            if intent:
                intent_p = doc.add_paragraph()
                intent_label = intent_p.add_run("考察意图：")
                intent_label.bold = True
                intent_label.font.size = Pt(11)
                intent_text = intent_p.add_run(intent)
                intent_text.font.size = Pt(11)

            points = q.get('evaluation_points', [])
            if points:
                points_label_p = doc.add_paragraph()
                points_label_run = points_label_p.add_run("评估要点：")
                points_label_run.bold = True
                points_label_run.font.size = Pt(11)
                for point in points:
                    doc.add_paragraph(f"  □ {point}", style='List Bullet')

        section_index += 1

    doc.save(str(output_path))
    return FileResponse(path=str(output_path), filename=filename,
                       media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
